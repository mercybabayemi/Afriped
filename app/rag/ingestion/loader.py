"""Document loaders for PDF and DOCX curriculum files."""
from __future__ import annotations

from pathlib import Path
from typing import List

from langchain_core.documents import Document
from loguru import logger


def load_pdf(path: Path) -> List[Document]:
    """Load a PDF and return one Document per page."""
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.error("pypdf not installed. Run: pip install pypdf")
        return []

    docs: List[Document] = []
    try:
        reader = PdfReader(str(path))
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source_file": path.name, "page": i + 1},
                ))
    except Exception as exc:
        logger.error(f"Failed to load PDF {path}: {exc}")

    logger.info(f"Loaded {len(docs)} pages from {path.name}")
    return docs


def load_docx(path: Path) -> List[Document]:
    """Load a DOCX and return one Document per non-empty paragraph."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return []

    docs: List[Document] = []
    try:
        docx = DocxDocument(str(path))
        current_text: List[str] = []
        for para in docx.paragraphs:
            text = para.text.strip()
            if text:
                current_text.append(text)
            elif current_text:
                # Save accumulated paragraph block
                docs.append(Document(
                    page_content="\n".join(current_text),
                    metadata={"source_file": path.name},
                ))
                current_text = []
        if current_text:
            docs.append(Document(
                page_content="\n".join(current_text),
                metadata={"source_file": path.name},
            ))
    except Exception as exc:
        logger.error(f"Failed to load DOCX {path}: {exc}")

    logger.info(f"Loaded {len(docs)} paragraphs from {path.name}")
    return docs


def load_file(path: Path) -> List[Document]:
    """Dispatch to appropriate loader based on file extension."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(path)
    elif suffix in {".docx", ".doc"}:
        return load_docx(path)
    elif suffix == ".txt":
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
            if text.strip():
                return [Document(
                    page_content=text,
                    metadata={"source_file": path.name},
                )]
        except Exception as exc:
            logger.error(f"Failed to load TXT {path}: {exc}")
        return []
    else:
        logger.warning(f"Unsupported file type: {suffix} ({path.name})")
        return []


def load_directory(directory: Path, recursive: bool = True) -> List[Document]:
    """Load all supported documents from a directory."""
    pattern = "**/*" if recursive else "*"
    supported = {".pdf", ".docx", ".doc", ".txt"}
    docs: List[Document] = []
    for fp in sorted(directory.glob(pattern)):
        if fp.is_file() and fp.suffix.lower() in supported:
            docs.extend(load_file(fp))
    logger.info(f"Total documents loaded from {directory}: {len(docs)}")
    return docs
